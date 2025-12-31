from flask import send_file
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
import io


def export_sales_pdf(data):
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    y = 750

    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(50, y, "Sales Report")
    y -= 30

    pdf.setFont("Helvetica", 10)

    for row in data:
        pdf.drawString(
            50, y,
            f"{row['period']} | Orders: {row['total_orders']} | Sales: ₱{row['total_sales'] or 0:.2f}"
        )
        y -= 20

        if y < 50:  # page break
            pdf.showPage()
            y = 750
            pdf.setFont("Helvetica", 10)

    pdf.save()
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="sales_report.pdf",
        mimetype="application/pdf"
    )


def export_sales_docx(data):
    buffer = io.BytesIO()
    doc = Document()
    doc.add_heading("Sales Report", level=1)

    for row in data:
        doc.add_paragraph(
            f"{row['period']} | Orders: {row['total_orders']} | Sales: ₱{row['total_sales'] or 0:.2f}"
        )

    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="sales_report.docx"
    )
